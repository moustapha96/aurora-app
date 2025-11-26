#!/usr/bin/env python3
"""
Script pour convertir AUDIT_ET_AMELIORATIONS.md en fichier Word (.docx)
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def read_markdown_file(filename):
    """Lit le fichier markdown"""
    with open(filename, 'r', encoding='utf-8') as f:
        return f.read()

def parse_markdown_to_docx(md_content, output_filename):
    """Convertit le contenu markdown en document Word"""
    doc = Document()
    
    # Style par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Titre principal
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # Sous-titre niveau 2
        if line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue
        
        # Sous-titre niveau 3
        if line.startswith('### '):
            text = line[4:].strip()
            # Supprimer les emojis pour le titre
            text = re.sub(r'[🔴🟠⚠️📄🔧📋✨🐛✅🔥⚡📅📆]', '', text).strip()
            doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Liste à puces
        if line.startswith('- '):
            text = line[2:].strip()
            # Supprimer les emojis de checklist
            text = re.sub(r'\[ \]|\[x\]', '', text).strip()
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Liste numérotée
        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Code block
        if line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
            i += 1
            continue
        
        # Ligne vide
        if not line:
            i += 1
            continue
        
        # Séparateur
        if line.startswith('---'):
            i += 1
            continue
        
        # Texte normal avec formatage
        p = doc.add_paragraph()
        parse_formatted_text(line, p)
        i += 1
    
    # Sauvegarder le document
    doc.save(output_filename)
    print(f"Document Word cree : {output_filename}")

def parse_formatted_text(text, paragraph):
    """Parse le texte avec formatage markdown"""
    # Pattern pour **gras**
    bold_pattern = r'\*\*(.*?)\*\*'
    # Pattern pour `code`
    code_pattern = r'`([^`]+)`'
    # Pattern pour les liens [texte](url)
    link_pattern = r'\[([^\]]+)\]\([^\)]+\)'
    
    # Remplacer les liens par juste le texte
    text = re.sub(link_pattern, r'\1', text)
    
    # Traiter le texte
    parts = []
    last_end = 0
    
    # Trouver tous les patterns
    matches = []
    for match in re.finditer(bold_pattern, text):
        matches.append(('bold', match.start(), match.end(), match.group(1)))
    for match in re.finditer(code_pattern, text):
        matches.append(('code', match.start(), match.end(), match.group(1)))
    
    # Trier par position
    matches.sort(key=lambda x: x[1])
    
    # Construire les runs
    for match_type, start, end, content in matches:
        # Ajouter le texte avant
        if start > last_end:
            paragraph.add_run(text[last_end:start])
        
        # Ajouter le texte formaté
        run = paragraph.add_run(content)
        if match_type == 'bold':
            run.bold = True
        elif match_type == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        
        last_end = end
    
    # Ajouter le reste
    if last_end < len(text):
        paragraph.add_run(text[last_end:])

if __name__ == '__main__':
    try:
        md_content = read_markdown_file('AUDIT_ET_AMELIORATIONS.md')
        parse_markdown_to_docx(md_content, 'AUDIT_ET_AMELIORATIONS.docx')
        print("\nConversion terminee avec succes!")
    except FileNotFoundError:
        print("Erreur: Fichier AUDIT_ET_AMELIORATIONS.md introuvable")
    except ImportError:
        print("Erreur: python-docx n'est pas installe")
        print("   Installez-le avec: pip install python-docx")
    except Exception as e:
        print(f"Erreur: {e}")

